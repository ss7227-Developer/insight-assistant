"""
document_loader.py
Loads PDF, DOCX, CSV/XLSX, JSON, and TXT files into LangChain Documents.
Each document gets metadata: source, domain, file_type, page_num, ingested_at.
"""

import os
import json
import hashlib
from datetime import datetime, timezone

import pandas as pd
from langchain_core.documents import Document


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".csv", ".xlsx", ".json"}


def load_documents(directory: str, domain: str) -> list[Document]:
    """
    Recursively loads all supported files from `directory` and returns
    a list of LangChain Documents tagged with the given domain.
    """
    docs: list[Document] = []
    for root, _, files in os.walk(directory):
        for filename in files:
            ext = os.path.splitext(filename)[1].lower()
            if ext not in SUPPORTED_EXTENSIONS:
                continue
            filepath = os.path.join(root, filename)
            try:
                loaded = _load_file(filepath, filename, domain, ext)
                docs.extend(loaded)
            except Exception as e:
                print(f"[document_loader] Skipping {filename}: {e}")
    return docs


def _load_file(filepath: str, filename: str, domain: str, ext: str) -> list[Document]:
    base_meta = {
        "source": filename,
        "domain": domain,
        "file_type": ext.lstrip("."),
        "ingested_at": datetime.now(timezone.utc).isoformat(),
    }

    if ext in (".txt", ".md"):
        return _load_text(filepath, base_meta)
    elif ext == ".pdf":
        return _load_pdf(filepath, base_meta)
    elif ext == ".docx":
        return _load_docx(filepath, base_meta)
    elif ext in (".csv", ".xlsx"):
        return _load_tabular(filepath, base_meta, ext)
    elif ext == ".json":
        return _load_json(filepath, base_meta)
    return []


def _load_text(filepath: str, meta: dict) -> list[Document]:
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        content = f.read()
    return [Document(page_content=content, metadata={**meta, "page_num": None})]


def _load_pdf(filepath: str, meta: dict) -> list[Document]:
    from pypdf import PdfReader

    reader = PdfReader(filepath)
    docs = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            docs.append(Document(
                page_content=text,
                metadata={**meta, "page_num": i + 1},
            ))
    return docs


def _load_docx(filepath: str, meta: dict) -> list[Document]:
    from docx import Document as DocxDocument

    doc = DocxDocument(filepath)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    content = "\n\n".join(paragraphs)
    return [Document(page_content=content, metadata={**meta, "page_num": None})]


def _load_tabular(filepath: str, meta: dict, ext: str) -> list[Document]:
    if ext == ".csv":
        df = pd.read_csv(filepath)
    else:
        df = pd.read_excel(filepath)

    docs = []
    cols = list(df.columns)
    for idx, row in df.iterrows():
        # Format each row as "Column: value | Column: value ..."
        row_text = " | ".join(f"{col}: {row[col]}" for col in cols)
        docs.append(Document(
            page_content=row_text,
            metadata={**meta, "page_num": None, "row_index": int(idx)},
        ))
    return docs


def _load_json(filepath: str, meta: dict) -> list[Document]:
    with open(filepath, "r", encoding="utf-8") as f:
        data = json.load(f)
    content = _flatten_json(data)
    return [Document(page_content=content, metadata={**meta, "page_num": None})]


def _flatten_json(obj, prefix: str = "") -> str:
    """Recursively flatten a JSON object into key: value lines."""
    lines = []
    if isinstance(obj, dict):
        for k, v in obj.items():
            key = f"{prefix}.{k}" if prefix else k
            lines.append(_flatten_json(v, prefix=key))
    elif isinstance(obj, list):
        for i, item in enumerate(obj):
            lines.append(_flatten_json(item, prefix=f"{prefix}[{i}]"))
    else:
        lines.append(f"{prefix}: {obj}")
    return "\n".join(lines)
